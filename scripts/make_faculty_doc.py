#!/usr/bin/env python
"""Generate a one-page faculty-facing brief for SERUM as a .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F, 0x35, 0x5E)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# --- base style ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    return p


# --- Title block ---
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("SERUM")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY
sub = doc.add_paragraph()
sr = sub.add_run("Content-Aware Containment of Malware When the Attack Is Unknown")
sr.italic = True
sr.font.size = Pt(12)
sr.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(2)

meta = doc.add_paragraph()
mr = meta.add_run("Amritha S.  •  Research project brief")
mr.font.size = Pt(9.5)
mr.font.color.rgb = GREY
meta.paragraph_format.space_after = Pt(8)

# --- One-line summary ---
oneliner = doc.add_paragraph()
ol = oneliner.add_run(
    "In one line: I am building a defender that figures out what a spreading "
    "computer worm is exploiting — without ever seeing the malware itself "
    "— and uses that to stop the outbreak far more efficiently than "
    "today’s methods."
)
ol.bold = True
ol.font.size = Pt(11)
oneliner.paragraph_format.space_after = Pt(8)

# --- The problem ---
heading("The problem I am tackling")
body(
    "A self-spreading worm can only infect a computer if that computer runs the "
    "specific vulnerable software the worm targets. So the network a worm "
    "actually travels through is not the whole network — it is only the "
    "smaller set of machines that share that one weakness."
)
body(
    "Today’s defenses ignore this. They protect the most-connected, "
    "“popular” machines in the network. But a highly connected machine "
    "is harmless if it cannot run the exploit — so protection spent on it is "
    "wasted. The real difficulty: the defender does not know which weakness the "
    "worm is attacking. It only sees which machines have fallen sick."
)

# --- The idea ---
heading("The key insight")
body(
    "Because a worm can only spread through vulnerable machines, every infected "
    "machine is a clue about which weakness is being exploited. My system reads "
    "these clues, builds a probabilistic belief about the hidden attack, and "
    "spends its limited defensive budget precisely on the machines that can "
    "actually catch it — defending cautiously at first, then sharply as the "
    "outbreak reveals its target."
)

# --- What I built ---
heading("What I have built")
bullet(
    "A simulator of a worm spreading through a realistic, mixed network of "
    "computers with different software."
)
bullet(
    "A smart “content-aware” defending agent that infers the unseen "
    "attack and allocates protection accordingly (a Bayesian / decision-under-"
    "uncertainty approach)."
)
bullet(
    "A mathematical result that says exactly when the hidden attack can be "
    "identified from an outbreak — confirmed correct 100% of the time in "
    "experiments."
)
bullet(
    "A full evaluation on real network data and real published software "
    "vulnerabilities (the NVD/CVE database)."
)

# --- Results ---
heading("Results so far")
body(
    "Against the strongest conventional defense, my agent reduces the number of "
    "infected machines to about 1.8% (versus 3.4%), while keeping far more "
    "machines online — and it does this without being told what the attack "
    "is, inferring it on its own. On a real organisational network the "
    "conventional defense barely helps, exactly as the theory predicts, and my "
    "method cuts infections by roughly 28%. All results use paired, "
    "statistically tested experiments."
)

# --- Why it matters ---
heading("Why it matters")
body(
    "It reframes malware containment around what is spreading, not just who is "
    "connected — and shows a defender can act intelligently against an attack "
    "it has never seen. The work sits at the intersection of network science, "
    "cyber-security, and AI decision-making, and is written up as a research "
    "paper."
)

# --- Status ---
status = doc.add_paragraph()
status.paragraph_format.space_before = Pt(8)
s1 = status.add_run("Current status:  ")
s1.bold = True
s1.font.color.rgb = NAVY
status.add_run(
    "Core system, theory, and experiments complete and tested; paper draft "
    "written. Next step is hardening the defense against a smarter, adaptive "
    "attacker."
)

doc.save("/Users/amritha/Desktop/SERUM_brief.docx")
print("saved -> /Users/amritha/Desktop/SERUM_brief.docx")
